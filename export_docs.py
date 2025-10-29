import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx não está instalado. Rode: pip install python-docx")
    raise


def add_paragraph(document: Document, text: str, style: str = None, bold: bool = False):
    p = document.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if style:
        p.style = style
    return p


def add_heading(document: Document, text: str, level: int):
    level = max(1, min(level, 4))
    document.add_heading(text, level=level)


def md_to_docx(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding='utf-8').splitlines()
    document = Document()

    # Fonte base
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    in_code_block = False
    bullet_buffer = []

    def flush_bullets():
        nonlocal bullet_buffer
        for item in bullet_buffer:
            p = document.add_paragraph(item)
            p.style = 'List Bullet'
        bullet_buffer = []

    for line in lines:
        if line.strip().startswith('```'):
            flush_bullets()
            in_code_block = not in_code_block
            if in_code_block:
                # inicia bloco de código (usaremos parágrafos monoespaçados)
                p = document.add_paragraph()
                p.style = 'No Spacing'
            continue

        if in_code_block:
            run = document.add_paragraph().add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            continue

        if line.startswith('# '):
            flush_bullets()
            add_heading(document, line[2:].strip(), 1)
        elif line.startswith('## '):
            flush_bullets()
            add_heading(document, line[3:].strip(), 2)
        elif line.startswith('### '):
            flush_bullets()
            add_heading(document, line[4:].strip(), 3)
        elif line.strip().startswith('- '):
            bullet_buffer.append(line.strip()[2:])
        elif line.strip() == '':
            flush_bullets()
            document.add_paragraph("")
        else:
            flush_bullets()
            add_paragraph(document, line)

    flush_bullets()

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(docx_path))
    # Evitar caracteres não suportados no console padrão do Windows
    print(f"Gerado: {docx_path}")


def main():
    targets = [
        ('RELATORIO_DEEP_LEARNING.md', 'docs/RELATORIO_DEEP_LEARNING.docx'),
        ('DOCUMENTACAO_PROJETO.md', 'docs/DOCUMENTACAO_PROJETO.docx'),
        ('RELATORIO_JUSTIFICATIVAS.md', 'docs/RELATORIO_JUSTIFICATIVAS.docx'),
    ]

    for src, dst in targets:
        src_path = Path(src)
        if not src_path.exists():
            print(f"⚠️ Arquivo não encontrado: {src}")
            continue
        md_to_docx(src_path, Path(dst))


if __name__ == '__main__':
    main()


